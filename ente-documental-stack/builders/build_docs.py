#!/usr/bin/env python3
"""
Builder: documentos .docx del stack (python-docx).

Genera con el estilo de casa MJM-FB (carátula + metadatos + historial de cambios).

Estructurales:
  - 00_INDICE_Stack_Documental_ENTE.docx           (índice DERIVADO del config: registro completo)
  - MJM-FB-TI-PLA-001-V0_Arquitectura_Stack_Documental_ENTE.docx   (esqueleto)
  - MJM-FB-TI-PLA-002-V0_Dataroom_Gobernanza.docx                  (esqueleto)
  - MJM-FB-TI-IT-001-V0_Dataroom_Deployment_Runbook.docx           (esqueleto)

PDD (andamiaje derivado del checklist maestro — build_checklist.SECTIONS):
  - MJM-FB-PR-INF-003-V0_Draft_PDD_ENTE.docx        (documento maestro con índice de sprints)
  - MJM-FB-PR-INF-004..011  (Sprints 1–8)  y  -012 (QA/QC)
    Cada requisito del checklist se vuelca como una sección a completar.

IMPORTANTE: son ANDAMIAJES para completar, NO documentos finales. No se fabrica el
contenido técnico del PDD; cada sección queda como encabezado + guía de qué cubrir.

Uso:
    python3 build_docs.py [--fecha AAAA-MM-DD] [--incluir estructurales|pdd|todos]
Salida:
    builders/output/*.docx
"""

import argparse
import datetime
import json
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(HERE, "..", "config", "dataroom.config.json")
OUT_DIR = os.path.join(HERE, "output")

# Reutiliza las secciones del checklist maestro como fuente del andamiaje del PDD.
sys.path.insert(0, HERE)
from build_checklist import SECTIONS  # noqa: E402

FOREST = RGBColor(0x1F, 0x4E, 0x3D)
GREY = RGBColor(0x70, 0x70, 0x70)
PROYECTO = "Proyecto Carbono ENTE — Río Negro (Ganadería Regenerativa)"
ESTANDAR = "Verra VCS + CCB"
PDD_FOLDER = "A_Expediente_Interno/01_PDD"


def load_config():
    with open(CONFIG_PATH, encoding="utf-8") as fh:
        return json.load(fh)


def parse_code(filename):
    stem = filename.rsplit(".", 1)[0]
    code, _, slug = stem.partition("_")
    titulo = slug.replace("_", " ").strip() or stem
    return code, titulo


def add_cover(doc, codigo, titulo, fecha, version="V0", clasificacion="Interno / Confidencial"):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(codigo)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = FOREST

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(titulo)
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = FOREST

    doc.add_paragraph()
    meta = [
        ("Código", codigo),
        ("Versión", version),
        ("Fecha", fecha),
        ("Proyecto", PROYECTO),
        ("Estándar", ESTANDAR),
        ("Clasificación", clasificacion),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(meta):
        table.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
        table.rows[i].cells[1].text = v
    doc.add_paragraph()


def add_change_history(doc, fecha, descripcion="Emisión inicial."):
    doc.add_heading("Historial de cambios", level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Light Grid Accent 1"
    for j, head in enumerate(["Versión", "Fecha", "Descripción", "Responsable"]):
        table.rows[0].cells[j].paragraphs[0].add_run(head).bold = True
    row = table.rows[1].cells
    row[0].text, row[1].text, row[2].text, row[3].text = "V0", fecha, descripcion, "Christian Farjat"


def add_outline(doc, sections):
    for heading, hint in sections:
        doc.add_heading(heading, level=1)
        p = doc.add_paragraph(hint)
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def new_doc(codigo, titulo, fecha):
    doc = Document()
    add_cover(doc, codigo, titulo, fecha)
    add_change_history(doc, fecha)
    doc.add_page_break()
    return doc


def build_indice(cfg, fecha):
    doc = new_doc("00_INDICE", "Índice del Stack Documental ENTE", fecha)
    doc.add_heading("Registro documental", level=1)
    doc.add_paragraph(
        "Índice generado automáticamente desde config/dataroom.config.json. "
        "Lista los entregables del expediente interno por carpeta destino."
    )
    for carpeta, files in cfg["fileMapping"].items():
        doc.add_heading(carpeta, level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for j, head in enumerate(["Código", "Título", "Archivo"]):
            table.rows[0].cells[j].paragraphs[0].add_run(head).bold = True
        for filename in files:
            code, titulo = parse_code(filename)
            cells = table.add_row().cells
            cells[0].text, cells[1].text, cells[2].text = code, titulo, filename

    # Backend (va como Google Sheet, no como binario en la carpeta).
    doc.add_heading("Backend (convertido a Google Sheet en Fase 3)", level=2)
    b = cfg["backend"]
    p = doc.add_paragraph()
    p.add_run("Fuente: ").bold = True
    p.add_run(b["sourceXlsx"])
    return "00_INDICE_Stack_Documental_ENTE.docx", doc


def build_arquitectura(fecha):
    doc = new_doc("MJM-FB-TI-PLA-001-V0", "Arquitectura del Stack Documental ENTE", fecha)
    add_outline(doc, [
        ("1. Propósito y alcance", "Objetivo del stack documental y qué cubre / qué no."),
        ("2. Estructura del dataroom", "Árbol de carpetas del Shared Drive y su racional (Expediente Interno, PDD, Base Habilitantes, Snapshots)."),
        ("3. Convención de códigos", "MJM-FB-<AREA>-<TIPO>-<NNN>-V<n>: significado de cada segmento."),
        ("4. Roles y niveles de acceso", "INTERNAL / VVB / BUYERS / AUDITOR y qué ve cada uno."),
        ("5. Flujo de aprobación y snapshots", "Del Draft al Approved for VVB; disparo del webhook y copia inmutable."),
        ("6. Backend AppSheet + Apps Script", "Tablas del backend, webhook y automatizaciones."),
        ("7. Seguridad y confidencialidad", "Repo privado, gitignore, manejo del WEBHOOK_TOKEN."),
        ("8. Anexos", "Diagramas, referencias, folder IDs."),
    ])
    return "MJM-FB-TI-PLA-001-V0_Arquitectura_Stack_Documental_ENTE.docx", doc


def build_gobernanza(fecha):
    doc = new_doc("MJM-FB-TI-PLA-002-V0", "Gobernanza del Dataroom", fecha)
    add_outline(doc, [
        ("1. Objeto", "Alcance de la gobernanza documental del dataroom."),
        ("2. Roles y responsabilidades", "RACI de los roles internos y externos (VVB, compradores, auditor)."),
        ("3. Matriz de accesos por rol", "Qué carpeta/documento puede ver/comentar/editar cada rol."),
        ("4. Ciclo de vida documental", "Estados (Draft, QA/QC, Approved for VVB, Shared, Verified) y transiciones."),
        ("5. Control de versiones", "Nomenclatura, versionado Vn y control de cambios."),
        ("6. Seguridad y confidencialidad", "Clasificación de la información y reglas de compartición externa."),
        ("7. Retención y archivo", "Política de retención y de snapshots inmutables."),
        ("8. Auditoría", "Uso de Audit_Log y trazabilidad de eventos."),
    ])
    return "MJM-FB-TI-PLA-002-V0_Dataroom_Gobernanza.docx", doc


def build_runbook_doc(fecha):
    doc = new_doc("MJM-FB-TI-IT-001-V0", "Dataroom — Deployment Runbook", fecha)
    add_outline(doc, [
        ("1. Precondiciones", "Cuenta Google, herramientas (git, python, clasp) y accesos."),
        ("2. Fase 0 — Entorno", "Verificación de herramientas y método de subida a Drive."),
        ("3. Fase 1 — Repo", "Versionado del stack (subcarpeta o repo propio)."),
        ("4. Fase 2 — Poblado del dataroom", "Subida de binarios y verificación de conteos (10 + 7)."),
        ("5. Fase 3 — Backend + Apps Script", "Conversión a Google Sheet, config y clasp push/deploy."),
        ("6. Pasos manuales", "Autorización de setupDataroom() y Bot de AppSheet."),
        ("Nota", "El detalle operativo vive en runbooks/DEPLOYMENT_RUNBOOK.md; este .docx es la versión de expediente."),
    ])
    return "MJM-FB-TI-IT-001-V0_Dataroom_Deployment_Runbook.docx", doc


# ── PDD (andamiaje derivado del checklist maestro) ────────────────────────────

def pdd_file_map(cfg):
    """Mapa prefijo-de-código -> (filename, code, titulo) para los archivos de 01_PDD."""
    out = {}
    for filename in cfg["fileMapping"][PDD_FOLDER]:
        code, titulo = parse_code(filename)      # code = MJM-FB-PR-INF-004-V0
        prefix = code.rsplit("-", 1)[0]          # MJM-FB-PR-INF-004
        out[prefix] = (filename, code, titulo)
    return out


def add_hint(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = GREY
    return p


def build_requisitos(doc, items):
    doc.add_heading("Requisitos a cubrir (del checklist maestro)", level=1)
    for requisito, ref, met_ref, obligatorio in items:
        doc.add_heading(requisito, level=2)
        add_hint(doc, f"Completar. Referencia: {ref} · Metodología: {met_ref} · Obligatorio: {obligatorio}.")


def build_pdd_doc(prefix, categoria, items, pmap, fecha):
    filename, code, titulo = pmap[prefix]
    doc = new_doc(code, titulo, fecha)
    doc.add_heading("1. Alcance y objetivo del sprint", level=1)
    add_hint(doc, f"Alcance: {categoria}. Completar objetivo, entradas, salidas y responsables del sprint.")
    build_requisitos(doc, items)
    return filename, doc


def build_pdd_master(pmap, fecha):
    filename, code, titulo = pmap["MJM-FB-PR-INF-003"]
    doc = new_doc(code, titulo, fecha)
    doc.add_heading("Estructura del PDD", level=1)
    doc.add_paragraph(
        "Documento maestro del PDD: consolida los sprints. Cada sprint tiene su documento propio "
        "(ver más abajo). El detalle de requisitos vive en el checklist maestro MJM-FB-PR-FOR-011."
    )
    for categoria, doc_ref, items in SECTIONS:
        doc.add_heading(categoria, level=2)
        sub = pmap.get(doc_ref)
        ref_txt = f"Documento: {sub[1]}" if sub else "Documento: (por asignar)"
        add_hint(doc, f"{ref_txt} · {len(items)} requisitos.")
    return filename, doc


def build_pdd_skeletons(cfg, fecha):
    pmap = pdd_file_map(cfg)
    docs = []
    if "MJM-FB-PR-INF-003" in pmap:
        docs.append(build_pdd_master(pmap, fecha))
    for categoria, doc_ref, items in SECTIONS:
        if doc_ref in pmap:
            docs.append(build_pdd_doc(doc_ref, categoria, items, pmap, fecha))
    return docs


def main():
    parser = argparse.ArgumentParser(description="Genera los .docx del stack ENTE.")
    parser.add_argument("--fecha", default=datetime.date.today().isoformat(),
                        help="Fecha de emisión AAAA-MM-DD (default: hoy).")
    parser.add_argument("--incluir", default="todos", choices=["estructurales", "pdd", "todos"],
                        help="Qué generar (default: todos).")
    args = parser.parse_args()

    cfg = load_config()
    os.makedirs(OUT_DIR, exist_ok=True)

    docs = []
    if args.incluir in ("estructurales", "todos"):
        docs += [
            build_indice(cfg, args.fecha),
            build_arquitectura(args.fecha),
            build_gobernanza(args.fecha),
            build_runbook_doc(args.fecha),
        ]
    if args.incluir in ("pdd", "todos"):
        docs += build_pdd_skeletons(cfg, args.fecha)

    for name, doc in docs:
        path = os.path.join(OUT_DIR, name)
        doc.save(path)
        print(f"OK  {path}")
    print(f"    {len(docs)} documentos generados · incluir={args.incluir} · fecha {args.fecha}")


if __name__ == "__main__":
    main()
